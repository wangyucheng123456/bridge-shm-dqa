"""Generate the MDPI Sensors cover letter as a Word .docx file."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_para(doc, text, bold=False, align=None, size=11, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def main() -> None:
    doc = Document()

    # MDPI A4 page setup matching the manuscript
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Header: sender block, top right
    add_para(doc, "[Given Name Surname]", align=WD_ALIGN_PARAGRAPH.RIGHT,
             space_after=2)
    add_para(doc, "[Department / School]", align=WD_ALIGN_PARAGRAPH.RIGHT,
             space_after=2)
    add_para(doc, "[University / Institution]",
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_para(doc, "[Street Address, City, Postal Code, Country]",
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_para(doc, "Email: [corresponding.author@inst.edu]",
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_para(doc, "Phone: [+xx-xxx-xxx-xxxx]",
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=14)

    # Date
    add_para(doc, "[Submission Date, e.g., 15 May 2026]",
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)

    # Recipient
    add_para(doc, "Editor-in-Chief", space_after=2)
    add_para(doc, "Sensors", bold=True, space_after=2)
    add_para(doc, "MDPI, St. Alban-Anlage 66, 4052 Basel, Switzerland",
             space_after=18)

    # Subject line
    add_para(doc,
             "Subject: Submission of an original research article to Sensors",
             bold=True, space_after=14)

    add_para(doc, "Dear Editor,", space_after=12)

    # Paragraph 1 -- what we submit
    add_para(doc,
        "On behalf of all co-authors I am pleased to submit our original "
        "research manuscript entitled \u201cData-Quality-Aware Damage "
        "Detection in Bridge Structural Health Monitoring: A Stress-Test "
        "Study on Two Real-World Bridges\u201d for consideration as a "
        "regular research article in Sensors. The submission package "
        "comprises the manuscript file Sensors_Bridge_SHM_DQA.docx, "
        "the figure files referenced therein, and this cover letter.",
        space_after=12)

    # Paragraph 2 -- what the paper is about
    add_para(doc,
        "Bridge structural health monitoring (SHM) increasingly relies on "
        "automated damage-sensitive analytics, but the value of those "
        "analytics is bounded by the integrity of the underlying sensor "
        "stream. Existing pipelines typically separate data-quality "
        "checks, denoising, and damage detection, and rarely measure "
        "their interaction on real bridge records with verified damage. "
        "Our work closes that loop. We define a four-dimension data-"
        "quality assessment (DQA) layer (Completeness, Accuracy, "
        "Consistency and signal-to-noise ratio) that is auditable, "
        "computable on a single edge node, and traceable to specific "
        "physical defect classes; we then subject two unsupervised "
        "damage detectors (a linear PCA reconstruction baseline and a "
        "compact deep autoencoder) to a controlled stress test on two "
        "independent real bridges with verified structural change\u2014"
        "the V\u00e4nersborg Bridge in Sweden (2023 fracture event) and "
        "the Z-24 Bridge in Switzerland (progressive damage). All "
        "metrics are summarised over twenty random seeds for the PCA "
        "detector and ten for the autoencoder, with bootstrap 95% "
        "confidence intervals, paired tests, and effect-size reporting.",
        space_after=12)

    # Paragraph 3 -- main findings & contribution
    add_para(doc,
        "Three results stand out. First, broadband electromagnetic "
        "noise is the dominant operational threat on both bridges: "
        "moderate noise (SNR \u2264 20 dB) collapsed the V\u00e4nersborg "
        "PCA AUC from 0.724 to 0.496, and severe noise drove both "
        "detectors towards chance-level discrimination, with Z-24 "
        "reproducing the same ordering. Second, none of the six "
        "classical repair pipelines tested could restore noise-degraded "
        "discriminability, while moving-average smoothing and median "
        "filtering recovered most of the loss for missingness and "
        "impulsive contamination, respectively. We frame this not as a "
        "negative result but as an actionable separation that the DQA "
        "layer makes explicit at ingest time: software-side defects "
        "should be repaired in software, but broadband noise must be "
        "addressed in hardware (shielding, grounding, low-noise "
        "front-end, gain staging, sensor placement). Third, the "
        "rank stability of per-sensor composite scores across six "
        "weighting schemes (mean Spearman \u03c1 = 0.999) supports the "
        "framework\u2019s comparative use even when absolute scores "
        "depend on the weighting prior.",
        space_after=12)

    # Paragraph 4 -- fit to the journal scope
    add_para(doc,
        "We believe the manuscript fits Sensors particularly well. "
        "It addresses the journal\u2019s sensor-centric audience by "
        "(i) operating directly on accelerometer and strain-gauge data "
        "from real in-service bridges; (ii) coupling sensor-data "
        "quality to downstream decision metrics rather than treating "
        "the two as separable; and (iii) translating the empirical "
        "findings into hardware-aware deployment guidance, including "
        "edge-node DQA and sensor-placement implications, that is "
        "directly usable by SHM practitioners and instrumentation "
        "engineers.",
        space_after=12)

    # Paragraph 5 -- declarations
    add_para(doc,
        "We confirm that this manuscript is the original work of the "
        "named authors, has not been published elsewhere in whole or "
        "in part, and is not under consideration by any other journal. "
        "All authors have read and approved the submitted version and "
        "agree to be accountable for the content. We declare no "
        "competing financial or non-financial interests that could be "
        "perceived to influence the work reported. The bridge "
        "measurement data used in this study are openly available "
        "under the licences specified in the original Zenodo and "
        "HuggingFace releases (DOI 10.5281/zenodo.8300495 for "
        "V\u00e4nersborg and the HuggingFace Hub release for Z-24), "
        "and all analysis scripts, controlled-degradation modules, "
        "repair operators and CSV result tables that reproduce the "
        "manuscript figures and tables are available from the "
        "corresponding author upon reasonable request.",
        space_after=12)

    # Paragraph 6 -- suggested reviewers (optional, helpful)
    add_para(doc,
        "Suggested reviewers (with no conflict of interest known to "
        "the authors) include researchers active in bridge SHM, "
        "data-quality assessment for sensor networks, and "
        "unsupervised damage detection. We will gladly provide "
        "specific names and affiliations on request.",
        space_after=12)

    # Closing
    add_para(doc,
        "We thank the editorial team and the reviewers in advance "
        "for their time and constructive feedback, and we look "
        "forward to your decision.",
        space_after=18)

    add_para(doc, "Sincerely,", space_after=2)
    add_para(doc, "[Given Name Surname]", space_after=2)
    add_para(doc, "Corresponding author", space_after=2)
    add_para(doc, "On behalf of all co-authors", space_after=0)

    out_path = "Cover_Letter_Sensors.docx"
    doc.save(out_path)
    print(f"[DONE] Cover letter saved to {out_path}")


if __name__ == "__main__":
    main()
